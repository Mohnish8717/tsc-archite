"""Layer 1: Contextual Ingestor (v4.0 — Concurrent SOTA).

Ingests raw input documents, normalizes, chunks semantically via LLM,
enriches with LLM-driven NLP, stores full data in Hindsight, and creates
a unified ProblemContextBundle.

v4.0 Changes:
  - Concurrent chunking: all documents chunked in parallel via asyncio.gather
  - Concurrent structured extraction: FeatureProposal + CompanyContext extracted
    in parallel using asyncio.gather, with LLM-fallback for non-JSON inputs
  - Enrichment semaphore raised 2→5 to saturate LeakyBucketQueue throughput
  - Structured timing, token, and rate-limit logs at every async boundary
"""

from __future__ import annotations

import csv
import io
import json
import logging
import os
import re
import time
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Optional



from tsc.llm.base import LLMClient
from tsc.llm.temperatures import L1_ENTITY_EXTRACTION, L1_PROBLEM_SYNTHESIS
from tsc.llm.prompts import (
    ENRICHMENT_SYSTEM,
    ENRICHMENT_USER,
    SEMANTIC_CHUNKING_SYSTEM,
    SEMANTIC_CHUNKING_USER,
)
from tsc.models.chunks import (
    ChunkEntity,
    EnrichedChunk,
    EntityType,
    ExtractedMetric,
    GlobalStatistics,
    ProblemContextBundle,
    SentimentLabel,
    SentimentResult,
    SourceSummary,
    TopicCategory,
)
from tsc.models.inputs import (
    CompanyContext,
    DocumentType,
    FeatureProposal,
    FileType,
    InputDocument,
    LoadedDocument,
    NormalizedContent,
)

logger = logging.getLogger(__name__)


# ── Custom Exceptions ────────────────────────────────────────────────


class ValidationError(Exception):
    """Raised when input validation fails."""


# ─────────────────────────────────────────────────────────────────────


class ContextualIngestor:
    """Layer 1: Load, normalize, chunk, enrich, and bundle input documents.

    v3.0: Stores full raw data in Hindsight before processing to prevent
    data loss. Uses LLM-only enrichment (spaCy removed).
    """

    def __init__(self, llm_client: LLMClient, session: Any = None) -> None:
        self._llm = llm_client
        self._session = session  # HindsightSessionManager for data retention
        logger.info("ContextualIngestor initialized (LLM-only enrichment)")

    # ── Public API ───────────────────────────────────────────────────

    async def process(
        self, documents: list[InputDocument]
    ) -> tuple[ProblemContextBundle, FeatureProposal, CompanyContext]:
        """Execute the full Layer 1 pipeline (v4.0 — concurrent).

        Returns:
            (ProblemContextBundle, FeatureProposal, CompanyContext)
        """
        import asyncio
        t0 = time.time()

        # Step 0: VALIDATE
        self._validate_inputs(documents)
        logger.info("⚡ L1 START — validating %d documents", len(documents))

        # Step 1.1: Load files (CPU-bound, sequential — file I/O is fast)
        loaded = [self._load_file(doc) for doc in documents]
        logger.info("✓ Loaded %d files in %.2fs", len(loaded), time.time() - t0)

        # Step 1.2: Normalize (CPU-bound, sequential)
        normalized = [self._normalize(doc) for doc in loaded]
        logger.info("✓ Normalized %d documents", len(normalized))

        # ── CONCURRENT: structured extraction + chunking + raw retention ──────
        # We fire all three in parallel because they are independent:
        #   a) extract_feature_proposal_async — LLM call if no JSON, else fast
        #   b) extract_company_context_async  — LLM call if no JSON, else fast
        #   c) retain_raw_documents           — Hindsight storage I/O
        #   d) semantic_chunk_v2              — LLM calls per doc (concurrent inside)
        t_concurrent = time.time()
        logger.info("⚡ L1 CONCURRENT — starting extraction + chunking in parallel")

        feature, company, _, chunks = await asyncio.gather(
            self._extract_feature_proposal_async(normalized),
            self._extract_company_context_async(normalized),
            self._retain_raw_documents(normalized),
            self._semantic_chunk_v2(normalized),
        )

        logger.info(
            "✓ Concurrent phase done in %.2fs — feature=%r, company=%r, chunks=%d",
            time.time() - t_concurrent, feature.title, company.company_name, len(chunks),
        )

        # Step 1.4: LLM Enrichment (concurrent across chunks via semaphore)
        t_enrich = time.time()
        enriched = await self._enrich_chunks(chunks)
        logger.info("✓ Enriched %d chunks in %.2fs", len(enriched), time.time() - t_enrich)

        # Quality gate
        quality = self._validate_enrichment_quality(enriched)
        logger.info(
            "✓ Quality: %.1f%% entities, %.1f%% metrics, avg_conf=%.3f",
            quality.get("pct_chunks_with_entities", 0),
            quality.get("pct_chunks_with_metrics", 0),
            quality.get("avg_entity_confidence", 0),
        )

        # DATA PRESERVATION: store enriched chunks
        await self._retain_enriched_chunks(enriched)

        # Step 1.5: Bundle
        bundle = self._create_bundle(enriched, time.time() - t0)
        logger.info(
            "✓ L1 DONE — %d chunks, %d entities, total=%.1fs",
            bundle.statistics.total_chunks,
            bundle.statistics.unique_entities,
            time.time() - t0,
        )

        return bundle, feature, company

    # ── CRITICAL FIX #1: Input Validation ────────────────────────────

    def _validate_inputs(self, documents: list[InputDocument]) -> None:
        """Validate input documents before processing.

        Raises:
            ValidationError: If any validation check fails.
        """
        if not documents:
            raise ValidationError("No documents provided")

        if len(documents) < 1:
            raise ValidationError(
                "At least one input document is required"
            )

        types_present = {d.type for d in documents}

        # Feature proposal is now OPTIONAL — Feature Discovery layer handles it
        if DocumentType.FEATURE_PROPOSAL not in types_present:
            logger.info("No FEATURE_PROPOSAL provided — Feature Discovery will generate one")

        if DocumentType.COMPANY_CONTEXT not in types_present:
            logger.warning("No COMPANY_CONTEXT provided — using defaults")

        # At minimum, we need either customer data OR a feature proposal
        content_types = {
            DocumentType.INTERVIEWS,
            DocumentType.SUPPORT_TICKETS,
            DocumentType.ANALYTICS,
        }
        has_content = any(t in types_present for t in content_types)
        has_proposal = DocumentType.FEATURE_PROPOSAL in types_present
        if not has_content and not has_proposal:
            raise ValidationError(
                "Need at least one input: customer data "
                "(interviews, support_tickets, analytics) or a feature_proposal"
            )

        for doc in documents:
            path = Path(doc.file_path)
            if not path.exists():
                raise ValidationError(f"File not found: {doc.file_path}")

    # ── v3.0: Hindsight Data Retention ──────────────────────────────

    async def _retain_raw_documents(
        self, normalized: list[NormalizedContent]
    ) -> None:
        """Store complete raw documents in Hindsight BEFORE chunking.

        Ensures zero data loss — even if chunking drops content (small
        sentences, malformed paragraphs), the full text is always available
        for downstream layers to query.
        """
        if not self._session:
            return

        for norm in normalized:
            doc_type = norm.document_type.value
            full_text = norm.normalized_text
            if not full_text:
                continue

            # Store complete document text — NO truncation
            if hasattr(self._session, "ingest_document"):
                await self._session.ingest_document(
                    document_text=full_text,
                    document_name=f"{doc_type}_input.txt",
                    doc_type=doc_type,
                    run_id="global"
                )
            else:
                await self._session.retain("world", full_text, metadata={
                    "type": "full_document",
                    "document_type": doc_type,
                    "word_count": len(full_text.split()),
                })

            # Store structured JSON if available (company context, proposals)
            if norm.json_parsed:
                await self._session.retain(
                    "world",
                    json.dumps(norm.json_parsed, indent=2),
                    metadata={
                        "type": "structured_data",
                        "document_type": doc_type,
                    },
                )

        logger.info("✓ Retained %d raw documents in Hindsight", len(normalized))

    async def _retain_enriched_chunks(
        self, chunks: list[EnrichedChunk]
    ) -> None:
        """Store enriched chunks with full metadata in Hindsight.

        Unlike the previous approach (which only stored chunk.text),
        this preserves entities, sentiment, urgency, topics, and metrics
        so downstream layers can query enrichment results.
        """
        if not self._session:
            return

        for chunk in chunks:
            sent_label = (
                chunk.sentiment.label.value
                if hasattr(chunk.sentiment.label, "value")
                else str(chunk.sentiment.label)
            )
            topic_val = (
                chunk.topic_category.value
                if hasattr(chunk.topic_category, "value")
                else str(chunk.topic_category)
            )

            await self._session.retain("world", chunk.text, metadata={
                "type": "enriched_chunk",
                "chunk_id": chunk.chunk_id,
                "source_type": chunk.source_type,
                "sentiment": sent_label,
                "urgency": chunk.urgency,
                "topic": topic_val,
                "entity_count": len(chunk.entities),
                "metric_count": len(chunk.metrics),
                "entities": [
                    {"text": e.text, "type": e.type if isinstance(e.type, str) else e.type.value}
                    for e in chunk.entities[:10]
                ],
            })

        logger.info("✓ Retained %d enriched chunks in Hindsight", len(chunks))

    # ── CRITICAL FIX #3: Enrichment Quality Gates ────────────────────

    def _validate_enrichment_quality(
        self, chunks: list[EnrichedChunk]
    ) -> dict[str, Any]:
        """Check enrichment coverage and confidence, warn if below targets."""
        if not chunks:
            return {}

        stats: dict[str, Any] = {
            "total_chunks": len(chunks),
            "chunks_with_entities": 0,
            "chunks_with_metrics": 0,
            "chunks_with_sentiment": 0,
            "avg_entity_confidence": 0.0,
            "min_entity_confidence": 1.0,
        }

        confidences: list[float] = []

        for chunk in chunks:
            if chunk.entities:
                stats["chunks_with_entities"] += 1
                confs = [e.confidence for e in chunk.entities]
                confidences.extend(confs)
                stats["min_entity_confidence"] = min(
                    stats["min_entity_confidence"], min(confs)
                )

            if chunk.metrics:
                stats["chunks_with_metrics"] += 1

            if chunk.sentiment:
                stats["chunks_with_sentiment"] += 1

        if confidences:
            stats["avg_entity_confidence"] = round(
                sum(confidences) / len(confidences), 3
            )

        # Convert to percentages
        total = len(chunks)
        stats["pct_chunks_with_entities"] = round(
            100 * stats["chunks_with_entities"] / total, 1
        )
        stats["pct_chunks_with_metrics"] = round(
            100 * stats["chunks_with_metrics"] / total, 1
        )
        stats["pct_chunks_with_sentiment"] = round(
            100 * stats["chunks_with_sentiment"] / total, 1
        )

        # Warnings if below targets
        if stats["pct_chunks_with_entities"] < 60:
            logger.warning(
                "Low entity coverage: %.1f%% (target 60%%)",
                stats["pct_chunks_with_entities"],
            )

        if stats["pct_chunks_with_metrics"] < 20:
            logger.warning(
                "Low metric coverage: %.1f%% (target 20%%)",
                stats["pct_chunks_with_metrics"],
            )

        if stats["avg_entity_confidence"] < 0.70:
            logger.warning(
                "Low entity confidence: %.3f (target 0.70)",
                stats["avg_entity_confidence"],
            )

        return stats

    # ── Step 1.1: File Loading ───────────────────────────────────────

    def _load_file(self, doc: InputDocument) -> LoadedDocument:
        path = Path(doc.file_path)
        if not path.exists():
            raise FileNotFoundError(f"Input file not found: {path}")

        file_type = self._detect_file_type(path)
        content = ""
        json_parsed = None
        csv_rows = None

        if file_type in (FileType.TXT, FileType.MD):
            content = self._read_text(path)
        elif file_type == FileType.PDF:
            content = self._read_pdf(path)
        elif file_type == FileType.DOCX:
            content = self._read_docx(path)
        elif file_type == FileType.JSON:
            raw = path.read_text(encoding="utf-8")
            json_parsed = json.loads(raw)
            content = raw
        elif file_type == FileType.CSV:
            csv_rows = self._read_csv(path)
            content = "\n".join(
                ", ".join(f"{k}: {v}" for k, v in row.items()) for row in csv_rows
            )

        if not content.strip():
            raise ValueError(f"Empty file after loading: {path}")

        return LoadedDocument(
            file_path=str(path),
            document_type=doc.type,
            file_type=file_type,
            content=content,
            json_parsed=json_parsed,
            csv_rows=csv_rows,
            file_size_kb=path.stat().st_size / 1024,
        )

    def _detect_file_type(self, path: Path) -> FileType:
        ext = path.suffix.lower().lstrip(".")
        mapping = {
            "pdf": FileType.PDF,
            "txt": FileType.TXT,
            "md": FileType.MD,
            "docx": FileType.DOCX,
            "json": FileType.JSON,
            "csv": FileType.CSV,
        }
        ft = mapping.get(ext)
        if not ft:
            raise ValueError(f"Unsupported file type: {ext}")
        return ft

    def _read_text(self, path: Path) -> str:
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Could not read text file with any encoding: {path}")

    def _read_pdf(self, path: Path) -> str:
        try:
            import pdfplumber

            with pdfplumber.open(path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except ImportError:
            logger.warning("pdfplumber not available, trying PyPDF2")
            from PyPDF2 import PdfReader

            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

    def _read_docx(self, path: Path) -> str:
        from docx import Document

        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)

    def _read_csv(self, path: Path) -> list[dict[str, Any]]:
        text = self._read_text(path)
        reader = csv.DictReader(io.StringIO(text))
        return list(reader)

    # ── Step 1.2: Normalization ──────────────────────────────────────

    def _normalize(self, doc: LoadedDocument) -> NormalizedContent:
        text = doc.content
        applied: list[str] = []

        # Standardize newlines
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        applied.append("newline_standardization")

        # Remove problematic special characters (keep useful punctuation)
        text = re.sub(r"[^\w\s.,?!:;()\"\'\n/-]", " ", text)
        applied.append("special_char_removal")

        # Collapse multiple spaces
        text = re.sub(r"[ \t]+", " ", text)
        applied.append("space_collapsing")

        # Strip leading/trailing
        text = "\n".join(line.strip() for line in text.split("\n"))
        text = re.sub(r"\n{3,}", "\n\n", text)
        applied.append("whitespace_stripping")

        # Quality score based on content length and ratio of normal chars
        alnum_ratio = sum(c.isalnum() for c in text) / max(len(text), 1)
        quality = min(1.0, alnum_ratio * 1.5)

        return NormalizedContent(
            document_type=doc.document_type,
            file_type=doc.file_type,
            normalized_text=text.strip(),
            json_parsed=doc.json_parsed,
            csv_rows=doc.csv_rows,
            normalization_applied=applied,
            quality_score=round(quality, 2),
        )

    # ── Step 1.3: Semantic Chunking ──────────────────────────────────

    async def _semantic_chunk_v2(
        self,
        normalized: list[NormalizedContent],
    ) -> list[EnrichedChunk]:
        """SOTA-1: LLM-driven semantic chunking — all docs chunked concurrently.

        v4.0: Each document is chunked in a separate asyncio task so that
        LLM latency for doc-A does not block doc-B. Final chunks are
        merged in original document order to maintain determinism.
        """
        import asyncio

        async def _chunk_one(norm: NormalizedContent) -> list[EnrichedChunk]:
            """Chunk a single normalised document, returning a list of raw (unindexed) chunks."""
            if not norm.normalized_text:
                return []

            word_count = len(norm.normalized_text.split())
            if word_count < 2000:
                logger.info(
                    "⚡ Chunking %s (%d words) — small doc, using paragraph splitter",
                    norm.document_type.value, word_count,
                )
                return self._simple_chunk_fallback(norm)

            t_doc = time.time()
            logger.info(
                "⚡ Chunking %s (%d words) — LLM semantic chunker started",
                norm.document_type.value, word_count,
            )
            try:
                prompt = SEMANTIC_CHUNKING_USER.render(document_content=norm.normalized_text)
                response_text = await self._llm.generate(
                    system_prompt=SEMANTIC_CHUNKING_SYSTEM,
                    user_prompt=prompt,
                    temperature=L1_ENTITY_EXTRACTION,
                    max_tokens=8000,
                )
                response_text = response_text.strip()
                if response_text.startswith("```json"):
                    response_text = response_text[7:-3].strip()
                elif response_text.startswith("```"):
                    response_text = response_text[3:-3].strip()

                raw_chunks = self._llm._parse_json_response(response_text)
                if not isinstance(raw_chunks, list):
                    if isinstance(raw_chunks, dict):
                        for v in raw_chunks.values():
                            if isinstance(v, list):
                                raw_chunks = v
                                break
                    if not isinstance(raw_chunks, list):
                        raise ValueError(f"Expected JSON list, got {type(raw_chunks)}")

                out: list[EnrichedChunk] = []
                for rc in raw_chunks:
                    chunk = EnrichedChunk(
                        chunk_id=rc.get("id", "chunk_tmp"),
                        text=rc.get("text", ""),
                        source_file=norm.file_type.value,
                        source_type=norm.document_type.value,
                        sequence=0,  # re-indexed below after merge
                        metadata=rc.get("metadata", {}),
                    )
                    if "speaker" in rc:
                        chunk.speaker_name = rc["speaker"]
                    elif "speaker" in chunk.metadata:
                        chunk.speaker_name = chunk.metadata["speaker"]
                    if "metadata" in rc and "primary_topic" in rc["metadata"]:
                        chunk.topics = [rc["metadata"]["primary_topic"]] + rc["metadata"].get("secondary_topics", [])
                    out.append(chunk)

                logger.info(
                    "✓ Chunked %s → %d chunks in %.2fs",
                    norm.document_type.value, len(out), time.time() - t_doc,
                )
                return out

            except Exception as e:
                logger.error(
                    "SOTA-1: chunking failed for %s (%.2fs elapsed), falling back: %s",
                    norm.document_type.value, time.time() - t_doc, e, exc_info=True,
                )
                return self._simple_chunk_fallback(norm)

        # Fire all documents concurrently; results come back in input order
        per_doc_chunks: list[list[EnrichedChunk]] = await asyncio.gather(
            *(_chunk_one(n) for n in normalized)
        )

        # Flatten and assign stable global indices
        all_chunks: list[EnrichedChunk] = []
        global_idx = 0
        for doc_chunks in per_doc_chunks:
            for c in doc_chunks:
                c.chunk_id = f"chunk_{global_idx:04d}"
                c.sequence = global_idx
                global_idx += 1
            all_chunks.extend(doc_chunks)

        return all_chunks

    def _simple_chunk_fallback(
        self, norm: NormalizedContent
    ) -> list[EnrichedChunk]:
        """Simple paragraph-based chunking fallback when LLM fails.

        Replaces the previous 115-line embedding-based chunker.
        Splits on double newlines and groups small paragraphs together.
        """
        chunks: list[EnrichedChunk] = []
        text = norm.normalized_text
        if not text:
            return chunks

        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        current_text: list[str] = []
        current_words = 0
        target_words = 800

        for para in paragraphs:
            para_words = len(para.split())

            if current_words + para_words > target_words and current_text:
                chunks.append(EnrichedChunk(
                    chunk_id=f"chunk_{len(chunks):04d}",
                    text="\n\n".join(current_text),
                    tokens=current_words,
                    source_file=norm.file_type.value,
                    source_type=norm.document_type.value,
                    sequence=len(chunks),
                ))
                current_text = [para]
                current_words = para_words
            else:
                current_text.append(para)
                current_words += para_words

        if current_text:
            chunks.append(EnrichedChunk(
                chunk_id=f"chunk_{len(chunks):04d}",
                text="\n\n".join(current_text),
                tokens=current_words,
                source_file=norm.file_type.value,
                source_type=norm.document_type.value,
                sequence=len(chunks),
            ))

        return chunks

    # ── Step 1.4: NLP Enrichment ─────────────────────────────────────

    async def _enrich_chunks(
        self, chunks: list[EnrichedChunk]
    ) -> list[EnrichedChunk]:
        """Enrich chunks with NER, sentiment, urgency, topics via LLM."""
        return await self._enrich_with_llm(chunks)

    async def _enrich_with_llm(
        self, chunks: list[EnrichedChunk]
    ) -> list[EnrichedChunk]:
        """Enrich using LLM — primary and only enrichment path (v4.0).

        Semaphore raised from 2→5 to saturate LeakyBucketQueue throughput.
        The rate limiter in gemini_provider/openai_provider still enforces
        TPM/RPM limits; the higher semaphore only removes artificial choking.
        """
        import asyncio
        sem = asyncio.Semaphore(5)
        
        async def enrich_chunk(chunk: EnrichedChunk):
            async with sem:
                # Segment-aware enrichment
                if chunk.source_type == "interviews":
                    chunk.is_customer_perspective = True
                elif chunk.source_type == "company_context":
                    chunk.is_customer_perspective = False

                try:
                    prompt = ENRICHMENT_USER.render(
                        text=chunk.text,
                        source_file=chunk.source_file,
                        source_type=chunk.source_type,
                    )
                    result = await self._llm.analyze(
                        system_prompt=ENRICHMENT_SYSTEM,
                        user_prompt=prompt,
                        temperature=L1_PROBLEM_SYNTHESIS,
                        max_tokens=1000,
                    )
                    self._apply_llm_enrichment(chunk, result)
                except Exception as e:
                    logger.warning("LLM enrichment failed for %s: %s", chunk.chunk_id, e)
                    # Minimal fallback: regex metrics + keyword urgency
                    chunk.metrics = self._extract_metrics(chunk.text)
                    chunk.urgency = self._estimate_urgency(chunk.text)

                # Hybrid pass: if general enrichment yielded 0 metrics, try dedicated LLM metric extraction
                if not chunk.metrics:
                    try:
                        llm_metrics = await self._extract_metrics_llm(chunk.text)
                        if llm_metrics:
                            chunk.metrics = llm_metrics
                            logger.info(
                                "LLM metric extraction recovered %d metrics for chunk %s",
                                len(llm_metrics), chunk.chunk_id,
                            )
                    except Exception as e:
                        logger.debug("LLM metric extraction failed for %s: %s", chunk.chunk_id, e)

                chunk.enrichment_timestamp = datetime.utcnow()

        await asyncio.gather(*(enrich_chunk(chunk) for chunk in chunks))
        return chunks

    def _apply_llm_enrichment(
        self, chunk: EnrichedChunk, result: dict[str, Any]
    ) -> None:
        """Apply LLM enrichment results to a chunk."""
        if not isinstance(result, dict):
            logger.warning(f"Expected dict for enrichment result, got {type(result)}: {result}")
            if isinstance(result, list) and len(result) > 0 and isinstance(result[0], dict):
                result = result[0]
            else:
                return

        # Entities
        raw_entities = result.get("entities", [])
        chunk.entities = [
            ChunkEntity(
                text=e.get("text", ""),
                type=e.get("type", "PRODUCT"),
                confidence=e.get("confidence", 0.5),
                value=e.get("value"),
                unit=e.get("unit"),
            )
            for e in raw_entities
        ]

        # Sentiment
        sent = result.get("sentiment", {})
        chunk.sentiment = SentimentResult(
            label=sent.get("label", "NEUTRAL"),
            score=sent.get("score", 0.5),
        )

        # Urgency
        chunk.urgency = result.get("urgency", 3)

        # Topic
        chunk.topic_category = result.get("topic_category", "feedback")
        chunk.topic_confidence = result.get("topic_confidence", 0.5)

        # Metrics — hybrid: merge LLM metrics with regex metrics, deduplicate
        raw_metrics = result.get("metrics", [])
        llm_metrics = [
            ExtractedMetric(
                value=m.get("value", 0),
                unit=m.get("unit", ""),
                context=m.get("context", ""),
            )
            for m in raw_metrics
        ]
        regex_metrics = self._extract_metrics(chunk.text)

        # Merge and deduplicate by value+unit signature
        seen_sigs: set[str] = set()
        merged: list[ExtractedMetric] = []
        for m in llm_metrics + regex_metrics:
            sig = f"{m.value}_{m.unit.lower()}"
            if sig not in seen_sigs:
                seen_sigs.add(sig)
                merged.append(m)
        chunk.metrics = merged


    def _extract_metrics(self, text: str) -> list[ExtractedMetric]:
        """Extract numeric metrics using expanded regex patterns."""
        metrics: list[ExtractedMetric] = []
        patterns = [
            # Original Patterns
            r"(\d+(?:\.\d+)?)\s*%\s*(?:of\s+)?(\w+)",
            r"(\d+(?:\.\d+)?)\s+(users?|customers?|tickets?|crashes?|requests?|times?)",
            r"\$(\d+(?:,\d+)*(?:\.\d+)?)\s*([KkMmB]?)",
            
            # Expanded Patterns
            # Temporal metrics (e.g., "20 minutes", "2 seconds")
            r"(\d+(?:\.\d+)?)\s*(minutes?|seconds?|hours?|days?|weeks?|months?|years?)",
            
            # Compound percentages (e.g., "10% reduction", "95% uptime")
            r"(\d+(?:\.\d+)?)\s*%\s*(reduction|increase|improvement|uptime|churn|adoption|growth|drop|fall)",
            
            # Multipliers (e.g., "3x faster", "10X improvement")
            r"(\d+(?:\.\d+)?)\s*[xX]\s*(faster|slower|improvement|growth|better|worse)",
            
            # Technical units (e.g., "500 MB", "20 ms")
            r"(\d+(?:\.\d+)?)\s*(MB|GB|TB|ms|fps|kbps|mbps)",
        ]
        
        extracted_signatures = set()
        
        for pattern in patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                try:
                    value_str = match.group(1).replace(",", "")
                    if not value_str:
                        continue
                        
                    value = float(value_str)
                    unit = match.group(2) if len(match.groups()) > 1 and match.group(2) else ""
                    
                    # Deduplicate exact matches
                    sig = f"{value}_{unit.lower()}"
                    if sig in extracted_signatures:
                        continue
                    extracted_signatures.add(sig)
                    
                    context = text[max(0, match.start() - 40) : min(len(text), match.end() + 40)]
                    metrics.append(
                        ExtractedMetric(value=value, unit=unit.strip(), context=context.strip())
                    )
                except (ValueError, IndexError):
                    pass
                    
        return metrics

    async def _extract_metrics_llm(self, text: str) -> list[ExtractedMetric]:
        """Dedicated LLM call to extract metrics when regex returns nothing.
        
        Uses a focused prompt that asks the LLM specifically for numeric facts,
        percentages, durations, counts, and monetary values embedded in the text.
        """
        prompt = f"""Extract ALL quantitative metrics, numbers, and measurements from this text.
Focus on:
- Percentages (e.g., "95% uptime", "10% churn")
- Durations (e.g., "20 minutes", "3 seconds latency")
- Counts (e.g., "500 users", "12 crashes per day")
- Money (e.g., "$50K ARR", "$2M revenue")
- Technical measurements (e.g., "200ms response", "4GB memory")
- Multipliers (e.g., "3x faster", "10x growth")
- Rates (e.g., "15 requests/second", "99.9% SLA")

Text:
---
{text[:1500]}
---

Return JSON:
{{
  "metrics": [
    {{"value": 95.0, "unit": "% uptime", "context": "brief surrounding text"}},
    {{"value": 200, "unit": "ms", "context": "response time under load"}}
  ]
}}

If no metrics are found, return {{"metrics": []}}.
Only return valid JSON, no markdown."""

        try:
            result = await self._llm.analyze(
                system_prompt="You are a precise data extraction specialist. Extract only metrics that are explicitly stated in the text. Do not infer or hallucinate numbers.",
                user_prompt=prompt,
                temperature=L1_ENTITY_EXTRACTION,
                max_tokens=800,
            )
            
            raw = result.get("metrics", [])
            metrics = []
            for m in raw:
                try:
                    val = float(m.get("value", 0))
                    unit = str(m.get("unit", ""))
                    ctx = str(m.get("context", ""))
                    if val != 0:
                        metrics.append(ExtractedMetric(value=val, unit=unit, context=ctx))
                except (ValueError, TypeError):
                    continue
                    
            logger.debug("LLM metric extraction found %d metrics", len(metrics))
            return metrics
            
        except Exception as e:
            logger.debug("LLM metric extraction error: %s", e)
            return []

    def _estimate_urgency(self, text: str) -> int:
        text_lower = text.lower()
        if any(
            w in text_lower
            for w in ("critical", "urgent", "blocking", "asap", "emergency")
        ):
            return 5
        if any(
            w in text_lower
            for w in ("important", "soon", "need", "must", "required")
        ):
            return 4
        if any(w in text_lower for w in ("would like", "nice to have", "want")):
            return 3
        if any(w in text_lower for w in ("maybe", "eventually", "consider")):
            return 2
        return 1


    # ── Step 1.5: Bundle Creation ────────────────────────────────────

    def _create_bundle(
        self, chunks: list[EnrichedChunk], processing_time: float
    ) -> ProblemContextBundle:
        """Create the unified ProblemContextBundle from enriched chunks."""
        # Build indices
        by_chunk_id: dict[str, Any] = {c.chunk_id: c.model_dump() for c in chunks}
        by_entity: dict[str, list[str]] = defaultdict(list)
        by_topic: dict[str, list[str]] = defaultdict(list)
        by_urgency: dict[str, list[str]] = defaultdict(list)
        by_sentiment: dict[str, list[str]] = defaultdict(list)

        entity_counter: Counter = Counter()
        topic_counter: Counter = Counter()
        sentiment_counter: Counter = Counter()
        urgency_sum = 0

        def get_val(v):
            return v.value if hasattr(v, "value") else str(v)

        for chunk in chunks:
            for ent in chunk.entities:
                by_entity[ent.text].append(chunk.chunk_id)
                entity_counter[ent.text] += 1

            tcat = get_val(chunk.topic_category)
            by_topic[tcat].append(chunk.chunk_id)
            topic_counter[tcat] += 1

            by_urgency[str(chunk.urgency)].append(chunk.chunk_id)

            slabel = get_val(chunk.sentiment.label)
            by_sentiment[slabel].append(chunk.chunk_id)
            sentiment_counter[slabel] += 1
            urgency_sum += chunk.urgency

        # Sources
        sources: dict[str, SourceSummary] = {}
        for chunk in chunks:
            st = chunk.source_type
            if st not in sources:
                sources[st] = SourceSummary()
            sources[st].count += 1
            sources[st].chunk_ids.append(chunk.chunk_id)

        return ProblemContextBundle(
            chunks=chunks,
            sources=sources,
            indices={
                "by_chunk_id": by_chunk_id,
                "by_entity": dict(by_entity),
                "by_topic": dict(by_topic),
                "by_urgency": dict(by_urgency),
                "by_sentiment": dict(by_sentiment),
            },
            statistics=GlobalStatistics(
                total_chunks=len(chunks),
                unique_entities=len(entity_counter),
                entity_summary=[
                    {"name": name, "mentions": count}
                    for name, count in entity_counter.most_common(20)
                ],
                topic_distribution=dict(topic_counter),
                sentiment_distribution=dict(sentiment_counter),
                average_urgency=round(urgency_sum / max(len(chunks), 1), 1),
            ),
            processing_stats={
                "total_files": len(sources),
                "total_chunks": len(chunks),
                "processing_time_seconds": round(processing_time, 1),
            },
        )

    # ── Structured Extraction ────────────────────────────────────────
    #
    # v4.0: Synchronous fast-path (JSON already parsed) preserved.
    # The async variants below add LLM-fallback for PDF/DOCX/TXT inputs.

    def _extract_feature_proposal(
        self, normalized: list[NormalizedContent]
    ) -> FeatureProposal:
        """Fast-path: return FeatureProposal when JSON was already parsed."""
        for n in normalized:
            if n.document_type == DocumentType.FEATURE_PROPOSAL and n.json_parsed:
                data = n.json_parsed
                return FeatureProposal(
                    title=data.get("title", "Unknown Feature"),
                    description=data.get("description", ""),
                    target_users=data.get("target_users", ""),
                    target_user_count=data.get("target_user_count"),
                    effort_weeks_min=data.get("effort_weeks_min")
                    or data.get("effort_weeks"),
                    effort_weeks_max=data.get("effort_weeks_max")
                    or data.get("effort_weeks"),
                    affected_domains=data.get("affected_domains", []),
                    existing_features=data.get("existing_features", []),
                    tech_stack=data.get("tech_stack", []),
                    priority=data.get("priority"),
                    revenue_model=data.get("revenue_model"),
                    pricing_strategy=data.get("pricing_strategy"),
                    customer_segments=data.get("customer_segments", []),
                )
        return FeatureProposal(
            title="Unspecified Feature", description="No proposal found"
        )

    async def _extract_feature_proposal_async(
        self, normalized: list[NormalizedContent]
    ) -> FeatureProposal:
        """Async extraction: JSON fast-path first, then LLM fallback for PDF/DOCX/TXT.

        The LLM fallback uses a strict JSON schema so the output always
        matches FeatureProposal fields — no post-processing guesswork.
        """
        # Fast-path: JSON already parsed (covers .json uploads)
        for n in normalized:
            if n.document_type == DocumentType.FEATURE_PROPOSAL and n.json_parsed:
                logger.info("📋 FeatureProposal — JSON fast-path (no LLM call needed)")
                return self._extract_feature_proposal(normalized)

        # LLM fallback: find raw text for FEATURE_PROPOSAL doc type
        for n in normalized:
            if n.document_type == DocumentType.FEATURE_PROPOSAL and n.normalized_text:
                t0 = time.time()
                logger.info(
                    "⚡ FeatureProposal — LLM extraction from %s text (%d chars)",
                    n.file_type.value, len(n.normalized_text),
                )
                system = (
                    "You are a product management analyst. Extract a structured feature "
                    "proposal from the provided document text. Return only valid JSON."
                )
                user = (
                    f"Extract a feature proposal from this document.\n\n"
                    f"---DOCUMENT---\n{n.normalized_text[:6000]}\n---END---\n\n"
                    "Return JSON with these fields (use null for missing values):\n"
                    '{"title": str, "description": str, "target_users": str, '
                    '"target_user_count": int|null, "effort_weeks_min": int|null, '
                    '"effort_weeks_max": int|null, "affected_domains": [str], '
                    '"existing_features": [str], "tech_stack": [str], '
                    '"priority": str|null, "revenue_model": str|null, '
                    '"pricing_strategy": str|null, "customer_segments": [str]}'
                )
                try:
                    data = await self._llm.analyze(
                        system_prompt=system,
                        user_prompt=user,
                        temperature=0.1,
                        max_tokens=1500,
                    )
                    logger.info(
                        "✓ FeatureProposal LLM extraction done in %.2fs — title=%r",
                        time.time() - t0, data.get("title"),
                    )
                    return FeatureProposal(
                        title=data.get("title") or "Unknown Feature",
                        description=data.get("description") or "",
                        target_users=data.get("target_users") or "",
                        target_user_count=data.get("target_user_count"),
                        effort_weeks_min=data.get("effort_weeks_min") or data.get("effort_weeks"),
                        effort_weeks_max=data.get("effort_weeks_max") or data.get("effort_weeks"),
                        affected_domains=data.get("affected_domains") or [],
                        existing_features=data.get("existing_features") or [],
                        tech_stack=data.get("tech_stack") or [],
                        priority=data.get("priority"),
                        revenue_model=data.get("revenue_model"),
                        pricing_strategy=data.get("pricing_strategy"),
                        customer_segments=data.get("customer_segments") or [],
                    )
                except Exception as e:
                    logger.warning(
                        "FeatureProposal LLM extraction failed (%.2fs): %s — using defaults",
                        time.time() - t0, e,
                    )

        logger.info("📋 FeatureProposal — no proposal document found, using defaults")
        return FeatureProposal(title="Unspecified Feature", description="No proposal found")

    def _extract_company_context(
        self, normalized: list[NormalizedContent]
    ) -> CompanyContext:
        """Fast-path: return CompanyContext when JSON was already parsed."""
        for n in normalized:
            if n.document_type == DocumentType.COMPANY_CONTEXT and n.json_parsed:
                data = n.json_parsed
                return CompanyContext(
                    company_name=data.get("company_name", ""),
                    team_size=data.get("team_size"),
                    budget=data.get("budget"),
                    tech_stack=data.get("tech_stack", []),
                    current_priorities=data.get("current_priorities", []),
                    competitors=data.get("competitors", []),
                    constraints=data.get("constraints", []),
                    stakeholders=data.get("stakeholders", []),
                )
        return CompanyContext()

    async def _extract_company_context_async(
        self, normalized: list[NormalizedContent]
    ) -> CompanyContext:
        """Async extraction: JSON fast-path first, then LLM fallback for PDF/DOCX/TXT."""
        # Fast-path
        for n in normalized:
            if n.document_type == DocumentType.COMPANY_CONTEXT and n.json_parsed:
                logger.info("📋 CompanyContext — JSON fast-path (no LLM call needed)")
                return self._extract_company_context(normalized)

        # LLM fallback: find raw text for COMPANY_CONTEXT doc type
        for n in normalized:
            if n.document_type == DocumentType.COMPANY_CONTEXT and n.normalized_text:
                t0 = time.time()
                logger.info(
                    "⚡ CompanyContext — LLM extraction from %s text (%d chars)",
                    n.file_type.value, len(n.normalized_text),
                )
                system = (
                    "You are a business analyst. Extract structured company context from "
                    "the provided document text. Return only valid JSON."
                )
                user = (
                    f"Extract company context from this document.\n\n"
                    f"---DOCUMENT---\n{n.normalized_text[:4000]}\n---END---\n\n"
                    "Return JSON with these fields (use null or [] for missing values):\n"
                    '{"company_name": str, "team_size": int|null, "budget": str|null, '
                    '"tech_stack": [str], "current_priorities": [str], '
                    '"competitors": [str], "constraints": [str], '
                    '"stakeholders": [{"name": str, "role": str}]}'
                )
                try:
                    data = await self._llm.analyze(
                        system_prompt=system,
                        user_prompt=user,
                        temperature=0.1,
                        max_tokens=1000,
                    )
                    logger.info(
                        "✓ CompanyContext LLM extraction done in %.2fs — company=%r",
                        time.time() - t0, data.get("company_name"),
                    )
                    return CompanyContext(
                        company_name=data.get("company_name") or "",
                        team_size=data.get("team_size"),
                        budget=data.get("budget"),
                        tech_stack=data.get("tech_stack") or [],
                        current_priorities=data.get("current_priorities") or [],
                        competitors=data.get("competitors") or [],
                        constraints=data.get("constraints") or [],
                        stakeholders=data.get("stakeholders") or [],
                    )
                except Exception as e:
                    logger.warning(
                        "CompanyContext LLM extraction failed (%.2fs): %s — using defaults",
                        time.time() - t0, e,
                    )

        logger.info("📋 CompanyContext — no context document found, using defaults")
        return CompanyContext()

